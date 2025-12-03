"""
Utility script to generate sample PDF and DOCX documents with dummy FAQs.

Run:
    python -m backend.sample_docs_generator
"""

from pathlib import Path

from docx import Document as DocxDocument
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas


def create_sample_docs(output_dir: str = "./sample_docs") -> None:
    base = Path(output_dir)
    base.mkdir(parents=True, exist_ok=True)

    # DOCX
    docx_path = base / "smartsupport_faq.docx"
    doc = DocxDocument()
    doc.add_heading("SmartSupport Cloud - FAQ", level=1)
    doc.add_paragraph("Q: What is SmartSupport Cloud?")
    doc.add_paragraph(
        "A: SmartSupport Cloud is a fictional AI-powered support platform used in this demo. "
        "It helps companies manage FAQs, automate responses, and provide knowledge search."
    )
    doc.add_paragraph("Q: How is my data stored?")
    doc.add_paragraph(
        "A: Data is stored securely in region-specific data centers with encryption at rest and in transit."
    )
    doc.add_paragraph("Q: Does SmartSupport Cloud support multi-agent workflows?")
    doc.add_paragraph(
        "A: Yes, SmartSupport Cloud can orchestrate multiple AI agents for classification, retrieval, and summarization."
    )
    doc.save(str(docx_path))

    # PDF
    pdf_path = base / "smartsupport_manual.pdf"
    c = canvas.Canvas(str(pdf_path), pagesize=letter)
    width, height = letter
    text = c.beginText(40, height - 40)
    text.setFont("Helvetica", 11)
    lines = [
        "SmartSupport Cloud - Onboarding Manual",
        "",
        "This manual explains how to onboard your team to SmartSupport Cloud.",
        "",
        "1. Create an account and invite your teammates.",
        "2. Upload your existing FAQ documents and product manuals.",
        "3. Configure your AI agents for general Q&A and document-based support.",
        "4. Connect SmartSupport Cloud to your helpdesk or chat widget.",
        "",
        "For security:",
        "- All connections use TLS.",
        "- Access control is role-based with audit logging.",
    ]
    for line in lines:
        text.textLine(line)
    c.drawText(text)
    c.showPage()
    c.save()

    print(f"Sample DOCX created at: {docx_path}")
    print(f"Sample PDF created at: {pdf_path}")


if __name__ == "__main__":
    create_sample_docs()


